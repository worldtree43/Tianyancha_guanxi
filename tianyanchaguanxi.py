from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO
import tkinter as tk
import webbrowser
from tkinter import filedialog
import time
import pandas as pd
import random
import os
import sys

def choose_excel_file():
    global excel_file
    file_path = filedialog.askopenfilename()
    if file_path:
        excel_file = file_path
        excel_label.config(text=f"已选择的 Excel 文件：{excel_file}")

def choose_downloads_folder():
    global downloads_folder
    folder_path = filedialog.askdirectory()
    if folder_path:
        downloads_folder = folder_path
        folder_label.config(text=f"已选择的下载文件夹：{downloads_folder}")


def show_instructions():
    instructions = """本程序通过selenium模拟搜索并批量下载天眼查公司关系图，效果如下：
            1. 需要用户手动登陆，预留时间为30秒。
            2. 用户需要选择其需要打开的 Excel 文件和默认的下载路径。
            3. 程序会自己打开 https://www.tianyancha.com/relation 网站，但需要用户自行登录，预设时间为30秒。
            4. 程序会自动读取用户选择的 Excel 文件中“Sheet1”中名为“公司列表”的列，对所有对象进行遍历，查询其两两之间的关系。
            5. 程序会下载公司的关系图，并将其整理放入一份名为“公司关联关系图”的 word 文档中，其位置为程序所处的文件夹。
            6. 如果您想了解更多信息，下载最新版本程序，或是获取程序源码，请访问程序的网站："""

    # 创建新的顶层窗口来显示使用说明
    top = tk.Toplevel(root)
    top.title("使用说明")
    tk.Label(top, text=instructions, justify=tk.LEFT, wraplength=400).pack(padx=10, pady=10)

    # 为网址创建链接，并在用户点击时用浏览器打开链接
    link_label = tk.Label(top, text="https://github.com/worldtree43/Tianyancha_guanxi", fg="blue", cursor="hand2", justify=tk.LEFT)
    link_label.pack()
    link_label.bind("<Button-1>", lambda event: webbrowser.open("https://github.com/worldtree43/Tianyancha_guanxi"))


def show_disclaimer():
    disclaimer = """免责声明：
    本程序仅供学习参考使用，任何人或组织不得将本仓库的内容用于非法用途或侵犯他人合法权益。
    本程序所涉及的爬虫技术仅用于学习和研究，不得用于对其他平台进行大规模爬虫或其他非法行为。
    对于因使用本程序内容而引起的任何法律责任，本程序不承担任何责任。
    使用本程序即表示您同意本免责声明的所有条款和条件。"""

    # 创建新的顶层窗口来显示免责声明
    top = tk.Toplevel(root)
    top.title("免责声明")
    tk.Label(top, text=disclaimer, justify=tk.LEFT).pack(padx=10, pady=10)


# 从excel中获取公司名称
def get_companies(file_path, sheet_name, column_name):
    df = pd.read_excel(file_path, sheet_name=sheet_name)
    return df[column_name].tolist()

def tianyancha_relation_screenshot(companies, download_folder):
    # 配置ChromeOptions，指定浏览器驱动路径
    chrome_options = Options()
    chrome_options.add_argument("--disable-gpu")  # 禁用GPU加速，解决部分环境下的问题
    chrome_options.add_argument("--window-size=1920x1080")  # 设置窗口大小

    # 创建一个Chrome浏览器实例，并最大化窗口
    driver = webdriver.Chrome(options=chrome_options)
    driver.maximize_window()

    # 打开天眼查关系页面
    url = 'https://www.tianyancha.com/relation'
    driver.get(url)

    # 设置一个等待时间，留出一定时间手动登录
    time.sleep(30)

    # 设置一个显式等待时间，等待搜索按钮可点击，最多等待10秒
    wait = WebDriverWait(driver, 10)

    for i in range(len(companies)):
        for j in range(i+1, len(companies)):
            try:
                search_input_from = wait.until(EC.presence_of_element_located((By.XPATH, '//*[@id="fromKey"]')))
                actions = ActionChains(driver)
                actions.move_to_element(search_input_from)
                actions.click()
                actions.send_keys(companies[i])
                actions.perform()

                search_input_to = wait.until(EC.presence_of_element_located((By.XPATH, '//*[@id="toKey"]')))
                actions = ActionChains(driver)
                actions.move_to_element(search_input_to)
                actions.click()
                actions.send_keys(companies[j])
                actions.perform()

                # 等待搜索按钮出现并点击
                search_button = wait.until(EC.element_to_be_clickable((By.XPATH, '//*[@id="web-content"]/section/div/div/div[1]/div[1]/span[2]')))
                driver.execute_script("arguments[0].click();", search_button)

                # 等待查询结果加载完成，这里设置为10秒，根据页面加载速度适当调整
                time.sleep(random.uniform(7, 10))

                # 等待图片加载完成点击下载按钮
                download_button = wait.until(EC.element_to_be_clickable((By.XPATH, '//*[@id="web-content"]/section/div/div/div[2]/div/div[3]/div/div[3]/div')))
                download_button.click()
                time.sleep(random.uniform(2, 5))

                close_button_from = wait.until(EC.element_to_be_clickable((By.XPATH,'//*[@id="web-content"]/section/div/div/div[1]/div[1]/div[1]/img')))
                close_button_from.click()
                close_button_to = wait.until(EC.element_to_be_clickable((By.XPATH, '//*[@id="web-content"]/section/div/div/div[1]/div[1]/div[3]/img')))
                close_button_to.click()

            except Exception as e:

                # 获取整个页面的截图
                screenshot = driver.get_screenshot_as_png()
                # 将截图数据转换成PIL的Image对象
                screenshot_image = Image.open(BytesIO(screenshot))

                # 设置保存截图的文件名，例如：'ab公司关联关系.png'
                filename = f"查关系图谱-{companies[i]}&{companies[j]}-天眼查.png"
                save_path = os.path.join(download_folder, filename)
                # 保存截图
                screenshot_image.save(save_path)
                driver.refresh()
                time.sleep(3)
    # 关闭浏览器
    driver.quit()

def insert_image_and_text(doc, image_path, text):
    try:
        doc.add_paragraph().add_run().add_picture(image_path, width=Inches(6.0))
        doc.add_paragraph(text)
    except FileNotFoundError as e:
        print(f"忽略图片 {image_path}，原因：{e}")

def create_word_document(images_info, downloads_folder):
    doc = Document()
    for image_info in images_info:
        image_name, text = image_info
        image_path = os.path.join(downloads_folder, image_name)
        insert_image_and_text(doc, image_path, text)
    doc.save("公司关联关系图.docx")

def process_files():
    sheet_name = 'Sheet1'
    company_column = '公司列表'
    images_info = []
    companies = get_companies(excel_file, sheet_name, company_column)
    tianyancha_relation_screenshot(companies, downloads_folder)

    for i in range(len(companies)):
        for j in range(i + 1, len(companies)):
            file_name = f"查关系图谱-{companies[i]}&{companies[j]}-天眼查.png"
            text = f"{companies[i]}&{companies[j]}关联关系\n"  # 换行
            image_info = (file_name, text)
            images_info.append(image_info)

    create_word_document(images_info,downloads_folder)
    sys.exit()

# 创建主窗口
root = tk.Tk()
root.title("天眼查关联关系图片批量下载程序")

button_frame = tk.Frame(root)
button_frame.pack(pady=10)

excel_file = ''
downloads_folder = ""

# 创建显示使用说明的按钮
instructions_btn = tk.Button(button_frame, text="使用说明", command=show_instructions)
instructions_btn.pack(side=tk.LEFT, padx=5)

# 创建显示免责声明的按钮
disclaimer_btn = tk.Button(button_frame, text="免责声明", command=show_disclaimer)
disclaimer_btn.pack(side=tk.RIGHT, padx=5)

# 创建选择文件和文件夹的按钮
excel_btn = tk.Button(root, text="选择Excel文件", command=choose_excel_file)
excel_btn.pack(pady=10)

folder_btn = tk.Button(root, text="选择下载文件夹", command=choose_downloads_folder)
folder_btn.pack(pady=10)

# 创建用于显示文件名称的标签
excel_label = tk.Label(root, text=f"已选择的 Excel 文件：{excel_file}")
excel_label.pack(pady=5)

folder_label = tk.Label(root, text=f"已选择的下载文件夹：{downloads_folder}")
folder_label.pack(pady=5)

# 创建处理文件的按钮
process_btn = tk.Button(root, text="开始下载图片并生成word文档", command=process_files)
process_btn.pack(pady=20)

# 开启主循环
root.mainloop()